from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
Space = " "


def Pics(document,File):
    pic = document.add_picture(File,width = Inches(3.86))
    Last = document.paragraphs[-1]
    Last.alignment = WD_ALIGN_PARAGRAPH.CENTER

def Line(Paragraph,Line,Size,Bold = False,NextL = True):
    if NextL == True :
        line0 = Paragraph.add_run("\n"+str(Line))
    elif NextL == False :
        line0 = Paragraph.add_run(str(Line))
    line0.bold = Bold
    line0_font = line0.font
    line0_font.size = Pt(Size)
    
def Line_sep(Var1,Seperator,Var2,Full,Var3):
    Spacing = Seperator - len(Var1)
    Spacing = Space*Spacing
    First_Part = Var1 + Spacing + Var2
    if Full == None:
        return First_Part
    elif Full == "SPECIAL":
        tolp = First_Part + str(Var3)
        return tolp
    else:
        x = Full - len(First_Part)
        Spacing =  x - len(Var3)
        Spacing = Space*Spacing
        Snd_Part = Spacing + Var3
        write = First_Part + Snd_Part
        return write
    
def Create_table(document):
    table = document.add_table(10,5)
    try :
        table.style = "Colorful List Accent 3"
    except(UserWarning):
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row0 = table.rows[0]
    row0.cells[0].text = "No"
    row0.cells[1].text = "Product"
    row0.cells[2].text = "Quantity"
    row0.cells[3].text = "Unit Price"
    row0.cells[4].text = "Amount"

    column0 = table.columns[0]
    for i in range(10):
        column0.cells[i].width = Inches(0.31)

    column1 = table.columns[1]
    for i in range(10):
        column1.cells[i].width = Inches(4.11)

    column2 = table.columns[2]
    for i in range(10):
        column2.cells[i].width = Inches(1.0)

    column3 = table.columns[3]
    for i in range(10):
        column3.cells[i].width = Inches(0.94)

    column4 = table.columns[4]
    for i in range(10):
        column4.cells[i].width = Inches(0.88)
    return table


def Assign_table(table,List):
    row = 0
    for i in List :
        row = row + 1
        rows1 = table.rows[row]
        for e in range(5):
            rows1.cells[e].text = str(i[e])
    
def Doc(Pic_Address,Company,Invoice_num,Invoice_Date,Company_Address,Customer_Name,Customer_Address,Email,Phone,Detail_top,Customer_Phone,Product_List,Extra_Info,Discount,Amount):
    document = Document("default.docx")

    section = document.sections
    for i in section :
        i.left_margin = Inches(0.75)
        i.right_margin = Inches(0.5)
        i.top_margin = Inches(0.4)
        i.bottom_margin = Inches(0.4)

    Pics(document,Pic_Address)
    para = document.add_paragraph("")
    Line1 = Line_sep(Var1 = Company,Seperator = 40,Var2 = "Invoice",Full = 60,Var3 = Invoice_num)
    Line(para,Line1,Size = 24,Bold = True,NextL = False)
    Line2 = Line_sep(Var1 = "",Seperator = 129,Var2 = "Invoice Date",Full = 161,Var3 = Invoice_Date)
    Line(para,Line2,Size = 11)
    Line(para,Company_Address,Size = 11)
    Line(para,Line = Detail_top,Size = 11)
    Line(para,Email,Size = 11)
    Line(para,Phone,Size = 11)
    #Customer
    Line3 = Line_sep(Var1 = "Bill To:",Seperator = 90,Var2 = "Ship To:",Full = None,Var3 = "")
    Line(para,Line3,Size = 11,Bold = True)
    Line(para,Line = Customer_Name,Size = 13,Bold = True)
    Line(para,Line = Customer_Address,Size = 11)
    Line(para,Line = Customer_Phone,Size = 11)
    #table
    table = Create_table(document)
    Assign_table(table,Product_List)
    #after table
    para1 = document.add_paragraph("")
    Line(para1,Line = Extra_Info,Size = 11,NextL = False)
    Sub_Total = Amount
    Line4 = Line_sep(Var1 = "",Seperator = 80,Var2 = "Sub Total     ",Full = "SPECIAL",Var3 = str(Sub_Total))
    Line(para1,Line = Line4,Bold = True,Size = 18)
    Line5 = Line_sep(Var1 = "",Seperator = 81,Var2 = "Discount     ",Full = "SPECIAL",Var3 = Discount)
    Line(para1,Line = Line5,Bold = True,Size = 18)
    Grand = float(Sub_Total)- float(Discount)
    Line6 = Line_sep(Var1 = "",Seperator = 57,Var2 = "Grand Total   ",Full = "SPECIAL",Var3 = str(Grand))
    Line(para1,Line = Line6,Bold = True,Size = 22)
    Line(para1,Line = " " ,Size = 12)
    line7 = "\n\nCustomer Signature                                                                                                                                     Signature"
    Line(para1,Line = line7 ,Size = 12)
    document.save("Invoice\Invoice   "+Invoice_num+".docx")

